"""Generate Word (.docx) reports on demand from live, role-scoped data."""

from __future__ import annotations

import io
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, RGBColor
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.models import AiUsage, Lesson, Progress, School, User
from app.models.enums import LessonStatus, Role, UserStatus
from app.services.ai_usage import (
    usage_breakdown_for_school,
    usage_by_user,
    usage_total_for_school,
)
from app.services.report_metrics import (
    QUIET_AFTER_DAYS,
    movement,
    progress_stats,
    quiet_teachers,
    security_anomalies,
)

BRAND = RGBColor(0x0F, 0x76, 0x6E)  # teal-700
MUTED = RGBColor(0x64, 0x74, 0x8B)  # slate-500


def _heading(doc: Document, text: str, size: int = 14) -> None:
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = BRAND


def _meta_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not rows:
        _meta_line(doc, "No records.")
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)


def _title_block(doc: Document, title: str, subtitle: str, generated_by: str) -> None:
    t = doc.add_paragraph()
    r = t.add_run("IM-Telligence")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = BRAND

    s = doc.add_paragraph()
    sr = s.add_run(title)
    sr.bold = True
    sr.font.size = Pt(16)

    when = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    _meta_line(doc, f"{subtitle}")
    _meta_line(doc, f"Generated {when} · by {generated_by}")
    doc.add_paragraph()


def _finish(doc: Document) -> io.BytesIO:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# --------------------------------------------------------------------------- #
def _school_sections(
    db: Session,
    doc: Document,
    school: School,
    include_security: bool = False,
    include_detail: bool = True,
) -> None:
    """One school's numbers.

    ``include_detail`` carries the per-assignment progress table. The platform
    report turns it off: one row per assignment per school is a document nobody
    opens, and the per-school report already exists for that.
    """
    teachers = list(
        db.scalars(
            select(User).where(User.school_id == school.id, User.role == Role.teacher)
        )
    )
    tids = [t.id for t in teachers]
    name_by_id = {t.id: t.name for t in teachers}
    progress = (
        list(db.scalars(select(Progress).where(Progress.teacher_id.in_(tids)))) if tids else []
    )
    lesson_ids = {p.lesson_id for p in progress}
    lessons = (
        {l.id: l.title for l in db.scalars(select(Lesson).where(Lesson.id.in_(lesson_ids)))}
        if lesson_ids
        else {}
    )

    usage = usage_by_user(db, tids)
    ai = usage_breakdown_for_school(db, school.id)

    stats = progress_stats(progress)
    moved = movement(db, tids, school_id=school.id)
    quiet = quiet_teachers(db, teachers)
    anomalies = security_anomalies(db, school_id=school.id) if include_security else []
    active = sum(1 for t in teachers if t.status == UserStatus.active)

    # Assigned / started / completed are three separate facts. The old single
    # "avg completion" averaged never-opened assignments in with the rest, so a
    # fresh upload made every school look like it had gone backwards.
    _heading(doc, "Summary")
    _table(
        doc,
        ["Active teachers", "Assigned", "Started", "Completed", "Not opened", "Late"],
        [
            [
                str(active),
                str(stats.assigned),
                str(stats.started),
                str(stats.completed),
                str(stats.not_started),
                str(stats.late),
            ]
        ],
    )
    _meta_line(
        doc,
        f"{stats.headline}."
        + (
            f" Among lessons actually begun, average progress is {stats.avg_of_started}%."
            if stats.avg_of_started is not None
            else ""
        ),
    )

    _heading(doc, "This week")
    for line in moved.lines():
        doc.add_paragraph(line, style="List Bullet")

    if quiet:
        _heading(doc, "Needs attention")
        _meta_line(
            doc, f"Active teachers with no lesson opened in {QUIET_AFTER_DAYS}+ days."
        )
        _table(
            doc,
            ["Teacher", "Email", "Last opened", "Completed to date"],
            [
                [
                    q.name,
                    q.email,
                    "never" if q.days_quiet is None else f"{q.days_quiet} days ago",
                    str(q.completed),
                ]
                for q in quiet
            ],
        )

    _heading(doc, "Teachers")
    _table(
        doc,
        ["Name", "Grades", "Language", "Status", "AI questions"],
        [
            [
                t.name,
                ", ".join(t.grades or []) or "—",
                (t.language or "—"),
                t.status.value,
                str(usage.get(t.id, {}).get("total", 0)),
            ]
            for t in teachers
        ],
    )

    _heading(doc, "AI assistant usage")
    _meta_line(doc, "Teacher questions to the lesson assistant — last 7 days and all time.")
    _meta_line(
        doc,
        f"School total {ai['total']}: {ai['teacher']} from teachers (below) + "
        f"{ai['admin']} from the school admin's operations assistant.",
    )
    _table(
        doc,
        ["Teacher", "Last 7 days", "Total questions"],
        [
            [
                t.name,
                str(usage.get(t.id, {}).get("last7", 0)),
                str(usage.get(t.id, {}).get("total", 0)),
            ]
            for t in sorted(
                teachers,
                key=lambda t: usage.get(t.id, {}).get("total", 0),
                reverse=True,
            )
        ],
    )

    if include_detail:
        # Late and unfinished work first — the top of the table is the part
        # anybody reads.
        def _urgency(p: Progress) -> tuple[int, str, str]:
            rank = {
                LessonStatus.late: 0,
                LessonStatus.in_progress: 1,
                LessonStatus.not_started: 2,
                LessonStatus.completed: 3,
            }
            return (rank.get(p.status, 4), name_by_id.get(p.teacher_id, ""), p.section)

        # A teacher who takes the same grade more than once has one row per
        # class, so the class has to be named or the table reads as duplicates
        # that contradict each other. Nobody else gets the column: for a teacher
        # with one class per grade it would be an empty column all the way down.
        sectioned = any(p.section for p in progress)

        _heading(doc, "Teacher progress")
        _meta_line(
            doc,
            "Late and in-progress lessons first."
            + (" One row per class." if sectioned else ""),
        )
        _table(
            doc,
            ["Teacher"]
            + (["Class"] if sectioned else [])
            + ["Lesson", "Status", "%", "Watchdog"],
            [
                [name_by_id.get(p.teacher_id, p.teacher_id)]
                + ([p.section or "—"] if sectioned else [])
                + [
                    lessons.get(p.lesson_id, p.lesson_id),
                    p.status.value,
                    str(p.percent_complete),
                    p.watchdog.value,
                ]
                for p in sorted(progress, key=_urgency)
            ],
        )

    if include_security:
        # Only what went wrong, grouped. Fifty lines of ordinary logins used to
        # bury the three events worth reading.
        _heading(doc, "Security alerts")
        if anomalies:
            _table(
                doc,
                ["User", "Event", "Status", "Times", "Last seen"],
                [
                    [
                        a.user_name,
                        a.event,
                        a.status,
                        str(a.count),
                        a.last_seen.strftime("%Y-%m-%d %H:%M") if a.last_seen else "—",
                    ]
                    for a in anomalies
                ],
            )
        else:
            _meta_line(doc, "No warnings or blocks in the last 30 days.")


def _clean_inline(text: str) -> str:
    # Strip the lightweight markdown emphasis the model tends to add.
    return text.replace("**", "").replace("__", "").strip()


def _render_narrative(doc: Document, narrative: str) -> None:
    """Render the assistant's markdown-ish narrative into Word paragraphs:
    `##`/`#` lines become headings, `-`/`*`/`•` lines become bullets, the rest
    are plain paragraphs."""
    for raw in narrative.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("#"):
            _heading(doc, _clean_inline(line.lstrip("#")), size=13)
        elif line[:2] in ("- ", "* ") or line.startswith("•"):
            doc.add_paragraph(_clean_inline(line.lstrip("-*• ")), style="List Bullet")
        else:
            p = doc.add_paragraph()
            run = p.add_run(_clean_inline(line))
            run.font.size = Pt(10)


def build_school_ai_report(
    db: Session, school_id: str, generated_by: str, narrative: str
) -> tuple[io.BytesIO, str]:
    """School report that leads with an AI-written executive narrative, followed
    by the same live data tables as the standard school report."""
    school = db.get(School, school_id)
    school_name = school.name if school else "School"
    doc = Document()
    _title_block(doc, "School Report", school_name, generated_by)

    # No wrapper heading: the narrative brings its own "## " headings, and a
    # section title above them only repeats the first one.
    _meta_line(doc, "Written from the live figures in the tables below.")
    _render_narrative(doc, narrative)
    doc.add_paragraph()

    if school:
        _school_sections(db, doc, school, include_security=False)
    date = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    return _finish(doc), f"IM-Telligence AI Report - {school_name} - {date}.docx"


def build_school_report(db: Session, school_id: str, generated_by: str) -> tuple[io.BytesIO, str]:
    school = db.get(School, school_id)
    school_name = school.name if school else "School"
    doc = Document()
    _title_block(doc, "School Report", school_name, generated_by)
    if school:
        _school_sections(db, doc, school, include_security=False)
    date = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    return _finish(doc), f"IM-Telligence Report - {school_name} - {date}.docx"


def _platform_body(db: Session, doc: Document) -> None:
    """Rollups and exceptions.

    This used to page-break into a full per-school dump, one row per assignment,
    which at curriculum scale produced a document nobody opened. Per-school
    detail lives in the per-school report; what belongs here is the comparison
    between schools and the handful of things that need somebody's attention.
    """
    schools = list(db.scalars(select(School).order_by(School.name)))
    teachers = list(db.scalars(select(User).where(User.role == Role.teacher)))
    lesson_count = db.scalar(select(func.count(Lesson.id))) or 0
    all_ids = [t.id for t in teachers]

    all_progress = (
        list(db.scalars(select(Progress).where(Progress.teacher_id.in_(all_ids))))
        if all_ids
        else []
    )
    stats = progress_stats(all_progress)
    moved = movement(db, all_ids)
    anomalies = security_anomalies(db)
    # Counted over the same 30 days the alerts table lists, so the summary and
    # the table can't tell the reader two different numbers.
    alerts_recent = sum(a.count for a in anomalies)
    ai_total = db.scalar(select(func.count(AiUsage.id))) or 0

    _heading(doc, "Platform summary")
    _table(
        doc,
        ["Schools", "Teachers", "Lessons", "Assigned"],
        [[str(len(schools)), str(len(teachers)), str(lesson_count), str(stats.assigned)]],
    )
    _table(
        doc,
        [
            "Started",
            "Completed",
            "Not opened",
            "Late",
            "Alerts (30d)",
            "AI questions (all time)",
        ],
        [
            [
                str(stats.started),
                str(stats.completed),
                str(stats.not_started),
                str(stats.late),
                str(alerts_recent),
                str(ai_total),
            ]
        ],
    )
    _meta_line(doc, f"{stats.headline}.")

    _heading(doc, "This week")
    for line in moved.lines():
        doc.add_paragraph(line, style="List Bullet")

    progress_by_school: dict[str, list[Progress]] = {s.id: [] for s in schools}
    school_of_teacher = {t.id: t.school_id for t in teachers}
    for p in all_progress:
        sid = school_of_teacher.get(p.teacher_id)
        if sid in progress_by_school:
            progress_by_school[sid].append(p)

    _heading(doc, "Schools overview")
    _meta_line(doc, "Completed out of assigned — a school with new uploads is not behind.")
    rows = []
    for s in schools:
        s_stats = progress_stats(progress_by_school.get(s.id, []))
        s_teachers = [t for t in teachers if t.school_id == s.id]
        rows.append(
            [
                s.name,
                s.city or "—",
                str(len(s_teachers)),
                str(s_stats.assigned),
                str(s_stats.completed),
                f"{s_stats.completion_rate}%",
                str(s_stats.late),
                str(usage_total_for_school(db, s.id)),
            ]
        )
    _table(
        doc,
        ["School", "City", "Teachers", "Assigned", "Done", "Rate", "Late", "AI"],
        rows,
    )

    # The exceptions, across every school, in one place.
    quiet = quiet_teachers(db, teachers)
    school_name_by_id = {s.id: s.name for s in schools}
    teacher_school = {t.email: school_name_by_id.get(t.school_id, "—") for t in teachers}
    _heading(doc, "Teachers needing a nudge")
    if quiet:
        _meta_line(doc, f"No lesson opened in {QUIET_AFTER_DAYS}+ days, quietest first.")
        _table(
            doc,
            ["Teacher", "School", "Last opened", "Completed to date"],
            [
                [
                    q.name,
                    teacher_school.get(q.email, "—"),
                    "never" if q.days_quiet is None else f"{q.days_quiet} days ago",
                    str(q.completed),
                ]
                for q in quiet
            ],
        )
    else:
        _meta_line(doc, "Every active teacher has opened a lesson recently.")

    _heading(doc, "Security alerts")
    if anomalies:
        _table(
            doc,
            ["User", "Event", "Status", "Times", "Last seen"],
            [
                [
                    a.user_name,
                    a.event,
                    a.status,
                    str(a.count),
                    a.last_seen.strftime("%Y-%m-%d %H:%M") if a.last_seen else "—",
                ]
                for a in anomalies
            ],
        )
    else:
        _meta_line(doc, "No warnings or blocks in the last 30 days.")


def build_super_report(
    db: Session, generated_by: str, school_id: str | None = None
) -> tuple[io.BytesIO, str]:
    date = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    # Scoped to one school -> same as the school report, with security.
    if school_id:
        school = db.get(School, school_id)
        school_name = school.name if school else "School"
        doc = Document()
        _title_block(doc, "School Report", school_name, generated_by)
        if school:
            _school_sections(db, doc, school, include_security=True)
        return _finish(doc), f"IM-Telligence Report - {school_name} - {date}.docx"

    doc = Document()
    _title_block(doc, "Platform Report", "All schools", generated_by)
    _platform_body(db, doc)
    return _finish(doc), f"IM-Telligence Platform Report - {date}.docx"


def build_super_ai_report(
    db: Session, generated_by: str, narrative: str
) -> tuple[io.BytesIO, str]:
    """Platform report that leads with the assistant's read of what needs
    attention, then the same rollups underneath it.

    The school admin has had this since the assistant shipped; the super-admin,
    who is the one deciding where to spend attention across every school, was
    getting tables alone.
    """
    date = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    doc = Document()
    _title_block(doc, "Platform Report", "All schools", generated_by)

    _meta_line(doc, "Written from the live figures in the tables below.")
    _render_narrative(doc, narrative)
    doc.add_paragraph()

    _platform_body(db, doc)
    return _finish(doc), f"IM-Telligence AI Platform Report - {date}.docx"
